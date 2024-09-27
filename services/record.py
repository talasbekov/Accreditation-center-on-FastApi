import os
from typing import Optional
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from io import BytesIO
from models import (
    Record,
)  # Предполагается, что у вас есть модель Record в models.py
from schemas import (
    RecordCreate,
    RecordUpdate,
)  # Предполагается, что у вас есть схемы создания и обновления событий
from services.base import ServiceBase


class RecordService(ServiceBase[Record, RecordCreate, RecordUpdate]):

    def get_by_name(self, db: Session, name: str) -> Optional[Record]:
        return db.query(Record).filter(Record.name == name).first()

    # Количество сотрудников по штату всего департамента
    def get_count_emp_by_state(self):
        return 136

    # Количество сотрудников по штату управления
    def get_count_state(self, db: Session):
        records = db.query(self.model).all()

        if records is None:
            return {"error": "Records not found"}

        recs = []
        for r in records:
            rec = db.query(Record).filter(Record.id == r.id).first()

            if rec is None:
                return {"error": "Record not found"}

            # Приведение всех значений к int с заменой None на 0
            count_state = int(rec.count_state or 0)
            count_list = int(rec.count_list or 0)
            count_on_leave = int(rec.count_on_leave or 0)
            count_on_sick_leave = int(rec.count_on_sick_leave or 0)
            count_business_trip = int(rec.count_business_trip or 0)
            count_seconded_in = int(rec.count_seconded_in or 0)
            count_seconded_out = int(rec.count_seconded_out or 0)
            count_on_duty = int(rec.count_on_duty or 0)
            count_after_on_duty = int(rec.count_after_on_duty or 0)
            count_at_the_competition = int(rec.count_at_the_competition or 0)

            # Рассчитываем количество вакантных мест и сотрудников "в строю"
            count_vacant = count_state - count_list
            count_in_service = count_list - (
                    count_on_leave
                    + count_on_sick_leave
                    + count_business_trip
                    + count_seconded_in
                    + count_seconded_out
                    + count_on_duty
                    + count_after_on_duty
                    + count_at_the_competition
            )

            recs.append( {
                "id": rec.id,
                "name": rec.name,
                "count_state": count_state,
                "count_list": count_list,
                "count_vacant": count_vacant,
                "count_in_service": count_in_service,
                "count_on_leave": count_on_leave,
                "count_on_sick_leave": count_on_sick_leave,
                "count_business_trip": count_business_trip,
                "count_seconded_in": count_seconded_in,
                "count_seconded_out": count_seconded_out,
                "count_on_duty": count_on_duty,
                "count_after_on_duty": count_after_on_duty,
                "count_at_the_competition": count_at_the_competition,
            }
            )
        return recs

    def create_word_report_from_template(self, recs: list):
        """
        Функция для создания отчета на основе шаблона Word с подстановкой данных
        """
        file_path = './test.docx'
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл {file_path} не найден. Пожалуйста, проверьте путь к файлу.")

        doc = Document(file_path)

        for table in doc.tables:
            row_idx = 5  # Начинаем с третьей строки

            for rec in recs:
                row = table.rows[row_idx]

                # Пример заполнения ячеек и применения форматирования
                for cell_idx, value in enumerate([
                    str(rec.get('id', '')),
                    rec.get('name', ''),
                    str(rec.get('count_state', 0)),
                    str(rec.get('count_list', 0)),
                    str(rec.get('count_in_service', 0)),
                    str(rec.get('count_vacant', 0)),
                    str(rec.get('count_on_leave', 0)),
                    str(rec.get('count_business_trip', 0)),
                    str(rec.get('count_on_sick_leave', 0)),
                    str(rec.get('count_on_duty', 0)),
                    str(rec.get('count_after_on_duty', 0)),
                    str(rec.get('count_at_the_competition', 0)),
                    str(rec.get('count_seconded_in', 0)),
                    str(rec.get('count_seconded_out', 0))
                ]):
                    cell = row.cells[cell_idx]
                    cell.text = value

                    # Применение выравнивания по центру (горизонтально и вертикально)
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Горизонтальное выравнивание

                    # Применение стиля к тексту
                    run = paragraph.runs[0]
                    run.font.size = Pt(12)  # Размер шрифта
                    run.font.name = 'Times New Roman'  # Имя шрифта

                    # Применение вертикального выравнивания
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    tcVAlign = OxmlElement('w:vAlign')
                    tcVAlign.set(qn('w:val'), 'center')
                    tcPr.append(tcVAlign)

                row_idx += 1

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    # def create_word_report_from_template(self, recs: list):
    #     """
    #     Функция для создания отчета на основе шаблона Word с подстановкой данных
    #     """
    #     file_path = './test.docx'
    #     # Проверка на существование файла
    #     if not os.path.exists(file_path):
    #         raise FileNotFoundError(f"Файл {file_path} не найден. Пожалуйста, проверьте путь к файлу.")
    #
    #     # Открываем заранее подготовленный шаблон Word
    #     doc = Document(file_path)
    #
    #     # Проходим по таблицам, начиная с третьей строки первой таблицы
    #     for table in doc.tables:
    #         row_idx = 5  # Начинаем с пятой строки (индексация начинается с 0)
    #
    #         for rec in recs:
    #             # Получаем строку таблицы для текущей записи
    #             row = table.rows[row_idx]
    #
    #             # Заполняем ячейки таблицы данными
    #             row.cells[0].text = str(rec.get('id', ''))
    #             row.cells[1].text = rec.get('name', '')
    #             row.cells[2].text = str(rec.get('count_state', 0))
    #             row.cells[3].text = str(rec.get('count_list', 0))
    #             row.cells[4].text = str(rec.get('count_in_service', 0))
    #             row.cells[5].text = str(rec.get('count_vacant', 0))
    #             row.cells[6].text = str(rec.get('count_on_leave', 0))
    #             row.cells[7].text = str(rec.get('count_business_trip', 0))
    #             row.cells[8].text = str(rec.get('count_on_sick_leave', 0))
    #             row.cells[9].text = str(rec.get('count_on_duty', 0))
    #             row.cells[10].text = str(rec.get('count_after_on_duty', 0))
    #             row.cells[11].text = str(rec.get('count_at_the_competition', 0))
    #             row.cells[12].text = str(rec.get('count_seconded_in', 0))
    #             row.cells[13].text = str(rec.get('count_seconded_out', 0))
    #
    #             # Переходим к следующей строке для следующей записи
    #             row_idx += 1
    #
    #     # Сохраняем измененный документ во временный буфер
    #     buffer = BytesIO()
    #     doc.save(buffer)
    #     buffer.seek(0)
    #
    #     # Возвращаем документ для выгрузки
    #     return buffer


record_service = RecordService(Record)
